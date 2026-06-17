from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "文档" / "AI用户购买品类偏好识别模块详细规格说明书.docx"

FONT = "Microsoft YaHei"
TITLE = "AI 用户购买品类偏好识别模块详细规格说明书"
SUBTITLE = "基于 BGE 中文语义向量模型的零售用户行为偏好分类"


def set_east_asia(run, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 14, 8),
        ("Heading 2", 13, "2E74B5", 10, 6),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_run(paragraph, text, bold=False, italic=False, size=None, color=None):
    run = paragraph.add_run(text)
    set_east_asia(run)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_para(doc, text="", style=None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True)
        add_run(p, text[len(bold_prefix) :])
    else:
        add_run(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_east_asia(run)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        set_cell_text(hdr[idx], text, bold=True, color="0B2545")
        shade_cell(hdr[idx], "E8EEF5")
        if widths:
            hdr[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], str(text))
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, body, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_run(p, title, bold=True, color="1F4D78")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, body)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F2F4F7")
    cell.text = ""
    for line in text.strip().splitlines():
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line.rstrip())
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(9)
    if cell.paragraphs and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]._element
        p.getparent().remove(p)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def set_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(header, "智能零售用户行为分析系统 | AI 模块规格说明", size=9, color="666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "AI 用户购买品类偏好识别模块", size=9, color="666666")


def build():
    doc = Document()
    style_doc(doc)
    set_header_footer(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, TITLE, bold=True, size=22, color="0B2545")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, SUBTITLE, size=12.5, color="4B5563")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p3, "适用项目：智能零售用户行为分析系统", size=11, color="4B5563")
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p4, "版本：V1.0    文档类型：详细规格说明书", size=10.5, color="4B5563")

    add_callout(
        doc,
        "文档摘要",
        "本说明书描述系统新增的 AI 用户购买品类偏好识别模块。模块引入 BAAI/bge-small-zh-v1.5 中文语义向量模型，"
        "将商品文本编码为语义向量，再结合用户浏览、点击、收藏、加购、下单、支付等行为权重，计算用户兴趣向量并识别其偏好商品品类。",
        fill="E8F4FF",
    )

    add_table(
        doc,
        ["项目", "内容"],
        [
            ["模块名称", "AI 用户购买品类偏好识别"],
            ["核心模型", "BAAI/bge-small-zh-v1.5，BGE 中文语义向量模型"],
            ["论文依据", "C-Pack: Packed Resources For General Chinese Embeddings，2023 年"],
            ["主要能力", "商品语义向量化、用户兴趣向量聚合、品类相似度分类"],
            ["前端入口", "ops001 登录后，左侧导航进入“AI偏好识别”"],
            ["后端接口", "GET /api/ai/preference-categories"],
        ],
        widths=[1.6, 4.9],
    )

    doc.add_page_break()

    add_para(doc, "1. 模块建设背景", "Heading 1")
    add_para(
        doc,
        "智能零售用户行为分析系统已经具备行为采集、指标分析、转化漏斗、商品热度、用户画像、智能推荐和报表统计等能力。"
        "在此基础上，新增 AI 模块的目标不是简单地用固定规则统计用户买过什么，而是引入真实预训练模型，让系统具备对商品文本语义的理解能力。"
    )
    add_para(
        doc,
        "传统统计方式只能看到“用户点击了某个分类”，但无法理解商品名称、品牌、标签、搜索词之间的语义关系。"
        "例如“轻量训练跑鞋”“城市复古板鞋”“运动缓震鞋”都属于相近消费兴趣，但字符串并不完全相同。"
        "BGE 向量模型可以把这些文本映射到语义空间，使相似商品在向量空间中更接近。"
    )
    add_callout(
        doc,
        "核心定位",
        "模型负责“看懂商品文字含义”，系统负责“结合用户行为判断偏好”。因此最终的用户偏好分类是模型能力与业务算法共同完成的结果。",
    )

    add_para(doc, "2. AI 模型与论文说明", "Heading 1")
    add_para(doc, "2.1 使用的模型", "Heading 2")
    add_para(
        doc,
        "本模块使用 BAAI/bge-small-zh-v1.5。BGE 是 BAAI General Embedding 的缩写，属于中文语义向量模型。"
        "该模型的输入是一段中文文本，输出是一组高维数字向量。向量中的数字并不是普通编号，而是模型对文本语义特征的表示。"
    )
    add_table(
        doc,
        ["模型信息", "说明"],
        [
            ["模型名称", "BAAI/bge-small-zh-v1.5"],
            ["模型类型", "中文语义向量模型 / Text Embedding Model"],
            ["主要用途", "文本语义匹配、检索、分类、聚类、相似度计算"],
            ["项目用途", "将商品名称、品牌、分类、标签、搜索词等文本编码为商品语义向量"],
            ["运行方式", "本地加载预训练权重，通过 sentence-transformers 完成向量推理"],
        ],
        widths=[1.8, 4.7],
    )

    add_para(doc, "2.2 论文通俗解释", "Heading 2")
    add_para(
        doc,
        "相关论文为《C-Pack: Packed Resources For General Chinese Embeddings》，2023 年发布。通俗地说，这篇论文的工作是为中文文本向量任务准备了一套资源包："
        "包括中文训练数据、中文评测基准，以及在这些资源上训练和评估的中文向量模型。"
    )
    add_para(
        doc,
        "论文关注的问题是：很多通用文本向量模型在英文上表现较好，但中文场景需要更适合中文语义、中文表达和中文任务的数据与模型。"
        "因此 C-Pack 通过构建中文语义任务资源，让模型能更好地理解中文文本之间的相似关系。"
    )
    add_bullet(doc, "如果两个商品文本含义相近，模型输出的向量距离会更近。")
    add_bullet(doc, "如果两个商品文本含义差异较大，模型输出的向量距离会更远。")
    add_bullet(doc, "系统可以利用这种“距离”来判断商品之间、用户兴趣与品类之间的相似度。")

    add_para(doc, "2.3 本项目实际使用了论文/模型的什么能力", "Heading 2")
    add_para(
        doc,
        "本项目没有重新训练模型，也没有使用模型直接输出“用户喜欢哪个分类”。实际使用的是 BGE 模型的文本向量化能力。"
        "也就是说，模型负责把商品文字转换成语义向量；系统再在这些向量基础上，结合用户行为权重完成偏好分类。"
    )
    add_table(
        doc,
        ["环节", "由谁完成", "说明"],
        [
            ["商品文本向量化", "BGE 预训练模型", "输入商品名称、品牌、分类、标签等文本，输出语义向量"],
            ["用户行为加权", "系统业务算法", "不同用户行为赋予不同权重，支付和加购权重大于浏览"],
            ["兴趣向量聚合", "系统业务算法", "将用户操作过的商品向量按权重加权求和"],
            ["品类偏好判断", "系统业务算法 + 向量相似度", "比较用户兴趣向量与品类向量，输出偏好品类和置信度"],
        ],
        widths=[1.5, 1.6, 3.4],
    )

    add_para(doc, "3. 功能规格说明", "Heading 1")
    add_para(doc, "3.1 功能目标", "Heading 2")
    add_bullet(doc, "识别用户最可能购买的商品品类，例如鞋服、食品、数码、家居等。")
    add_bullet(doc, "展示每个用户的主偏好品类、置信度、语义相似度、行为分和证据商品。")
    add_bullet(doc, "让运营人员可以理解用户为什么被判定为某一品类偏好。")
    add_bullet(doc, "为后续精准推荐、优惠券投放、用户分群和运营策略提供基础。")

    add_para(doc, "3.2 页面功能", "Heading 2")
    add_table(
        doc,
        ["页面区域", "展示内容", "用途"],
        [
            ["模型权重卡片", "已加载 / 待加载，模型名称 BAAI/bge-small-zh-v1.5", "确认是否使用真实预训练模型"],
            ["推理模式卡片", "语义向量 / 兜底向量", "说明当前是否由 sentence-transformers 执行向量推理"],
            ["识别用户卡片", "当前识别到的用户数量", "展示参与 AI 分类的用户样本数"],
            ["平均置信度卡片", "Top-1 品类平均置信度", "辅助判断模型结果整体可信度"],
            ["模型说明区", "模型名称、权重状态、生成时间、行为权重", "解释模型来源和权重配置"],
            ["用户偏好卡片", "用户 ID、主偏好品类、各候选品类、证据商品", "支持运营人员查看分类结果"],
        ],
        widths=[1.55, 2.7, 2.25],
    )

    add_para(doc, "3.3 输入与输出", "Heading 2")
    add_table(
        doc,
        ["类型", "字段", "说明"],
        [
            ["输入", "productName", "商品名称，如“轻量训练跑鞋 11”"],
            ["输入", "category", "商品类目，如“鞋服”“数码”"],
            ["输入", "brand", "品牌名称，用于补充商品语义"],
            ["输入", "tag", "商品标签，如“热卖”“高意向”等"],
            ["输入", "eventType", "用户行为类型，如 browse、click、cart_add、payment_success"],
            ["输出", "primaryCategory", "系统识别出的用户主偏好品类"],
            ["输出", "primaryConfidence", "主偏好品类置信度，范围 0-100"],
            ["输出", "evidenceProducts", "影响该判断的主要证据商品"],
        ],
        widths=[0.9, 1.7, 3.9],
    )

    add_para(doc, "4. 技术选型", "Heading 1")
    add_table(
        doc,
        ["技术项", "选型", "选择原因"],
        [
            ["AI 模型", "BAAI/bge-small-zh-v1.5", "中文语义向量模型，适合中文商品文本理解"],
            ["模型运行库", "sentence-transformers", "封装文本向量编码流程，调用简单，适合本地推理"],
            ["深度学习框架", "PyTorch CPU", "本地部署不依赖 GPU，适合实训演示环境"],
            ["模型下载源", "ModelScope", "国内下载稳定，已完成本地权重缓存"],
            ["后端框架", "FastAPI", "项目原有技术栈，便于新增 REST API"],
            ["前端框架", "React + TypeScript", "项目原有控制台架构，便于新增 AI 页面"],
            ["数据来源", "行为事件仓库", "复用现有用户行为采集链路，不新增复杂数据表"],
        ],
        widths=[1.4, 2.1, 3.0],
    )

    add_para(doc, "5. 算法与实现流程", "Heading 1")
    add_para(doc, "5.1 总体流程", "Heading 2")
    for step in [
        "采集用户行为事件，包含用户 ID、商品 ID、行为类型、发生时间和商品 metadata。",
        "将商品名称、品牌、分类、标签、搜索词、价格等字段拼接成商品语义文本。",
        "调用 BGE 模型，将每个商品语义文本编码为向量。",
        "根据行为类型设置权重，浏览权重较低，收藏、加购、支付权重较高。",
        "将用户操作过的商品向量按行为权重聚合，形成用户兴趣向量。",
        "按商品类别聚合类别向量，并计算用户兴趣向量与类别向量的余弦相似度。",
        "结合语义相似度和行为分，输出候选品类排序、主偏好品类和置信度。",
    ]:
        add_number(doc, step)

    add_para(doc, "5.2 行为权重设计", "Heading 2")
    add_table(
        doc,
        ["行为类型", "中文含义", "权重", "设计理由"],
        [
            ["browse", "浏览", "0.20", "弱兴趣信号，只能说明用户接触过商品"],
            ["search", "搜索", "0.30", "主动表达需求，但还未点击具体商品"],
            ["click", "点击", "0.40", "对商品产生初步兴趣"],
            ["favorite", "收藏", "0.65", "用户保存商品，兴趣强于普通点击"],
            ["cart_add", "加购", "0.85", "接近购买决策，属于强意向行为"],
            ["order_submit", "下单", "0.95", "已进入交易流程，强购买意图"],
            ["payment_success", "支付成功", "1.15", "已完成购买，是最强偏好信号"],
        ],
        widths=[1.25, 1.05, 0.75, 3.45],
    )

    add_para(doc, "5.3 评分公式说明", "Heading 2")
    add_callout(
        doc,
        "简化公式",
        "用户兴趣向量 = Σ(商品向量 × 行为权重 × 时间衰减 × 金额修正)。"
        "品类置信度 = 语义相似度得分 × 70% + 行为强度得分 × 30%。",
        fill="FFF8E8",
    )
    add_para(
        doc,
        "其中，语义相似度通过余弦相似度计算。行为强度得分来自该用户在某一类目下的加权行为累计值。"
        "这样设计可以同时体现“商品文本语义是否接近”和“用户行为是否足够强”。"
    )

    add_para(doc, "5.4 为什么不是简单规则统计", "Heading 2")
    add_para(
        doc,
        "如果只做规则统计，系统只能统计“某个分类出现多少次”。当前模块引入 BGE 后，商品文本本身会先被模型理解为语义向量。"
        "即使商品名称不完全一样，只要语义相近，也能在向量空间中形成较近的距离。"
        "因此本模块是“预训练语义模型 + 行为权重算法”的组合，而不是单纯的 if-else 或 SQL 计数。"
    )

    add_para(doc, "6. 后端具体实现", "Heading 1")
    add_table(
        doc,
        ["文件", "职责"],
        [
            ["backend/src/services/ai_preference_service.py", "AI 偏好识别核心服务，负责模型加载、向量编码、用户兴趣向量聚合和品类分类"],
            ["backend/src/api/routes/ai_preferences.py", "新增 FastAPI 路由，提供 /api/ai/preference-categories 接口"],
            ["backend/src/api/app.py", "注册 AI 路由，使接口纳入后端应用"],
            ["backend/scripts/download_ai_model.py", "模型下载脚本，用于提前缓存 BGE 预训练模型权重"],
            ["backend/tests/unit/test_ai_preference_service.py", "AI 服务单元测试，验证行为权重和分类结果"],
        ],
        widths=[2.7, 3.8],
    )

    add_para(doc, "6.1 模型加载策略", "Heading 2")
    add_para(
        doc,
        "服务启动后不会立刻加载大模型，而是在首次请求 AI 接口时懒加载模型。为避免并发请求重复加载模型，服务中加入了进程级单例和线程锁。"
        "模型加载完成后会复用同一个 SentenceTransformer 实例，后续请求直接进行向量编码。"
    )
    add_bullet(doc, "正常模式：加载 BGE 预训练权重，mode 为 sentence-transformers，loaded 为 true。")
    add_bullet(doc, "兜底模式：如果依赖或权重缺失，临时使用 deterministic hash embedding，保证系统页面不崩溃。")
    add_bullet(doc, "页面展示：只显示模型名称和权重状态，不暴露本地模型路径。")

    add_para(doc, "6.2 接口定义", "Heading 2")
    add_code(
        doc,
        """
GET /api/ai/preference-categories

Query:
  target_id: string = all
  top_n: integer = 5

Response:
  model: 模型信息
  behaviorWeights: 行为权重配置
  items: 用户偏好识别结果列表
  generatedAt: 结果生成时间
        """,
    )
    add_table(
        doc,
        ["返回字段", "含义"],
        [
            ["model.name", "模型名称，页面统一展示为 BAAI/bge-small-zh-v1.5"],
            ["model.mode", "推理模式，sentence-transformers 表示真实模型推理"],
            ["model.loaded", "模型权重是否加载成功"],
            ["items[].userId", "用户 ID"],
            ["items[].primaryCategory", "用户最可能偏好的商品品类"],
            ["items[].primaryConfidence", "主偏好品类置信度"],
            ["items[].categories", "候选品类列表，包含相似度、行为分、证据商品"],
        ],
        widths=[2.1, 4.4],
    )

    add_para(doc, "7. 前端具体实现", "Heading 1")
    add_table(
        doc,
        ["文件", "职责"],
        [
            ["frontend/src/services/aiPreferenceApi.ts", "封装 AI 偏好识别接口调用和 TypeScript 类型"],
            ["frontend/src/App.tsx", "新增 ViewKey、导航项、状态管理、AI 页面组件 AIPreferenceView"],
            ["frontend/src/styles.css", "新增 AI 页面卡片、权重标签、用户偏好卡片等样式"],
        ],
        widths=[2.6, 3.9],
    )
    add_para(
        doc,
        "前端页面重点突出“AI 模型已加载”“当前推理模式”“识别用户数量”“平均置信度”。"
        "为避免展示开发机器信息，页面不会显示本地模型路径，只展示模型名称、权重状态和业务解释。"
    )

    add_para(doc, "8. 部署与启动说明", "Heading 1")
    add_para(doc, "8.1 本地演示启动方式", "Heading 2")
    add_para(doc, "后端启动命令：")
    add_code(
        doc,
        r"""
cd E:\shixun\zy

$env:APP_ENV="development"
$env:USE_IN_MEMORY_REPOSITORY="1"
$env:AI_PREFERENCE_MODEL="E:\shixun\zy\var\ai_models\modelscope\BAAI\bge-small-zh-v1___5"
$env:AI_MODEL_CACHE_DIR="E:\shixun\zy\var\ai_models"

.\backend\.venv\Scripts\python.exe -m uvicorn api.app:app --host 127.0.0.1 --port 8000
        """,
    )
    add_para(doc, "前端启动命令：")
    add_code(
        doc,
        r"""
cd E:\shixun\zy\frontend
npm install
npm run dev -- --host 127.0.0.1 --port 5173
        """,
    )
    add_para(doc, "访问地址与账号：")
    add_table(
        doc,
        ["项目", "值"],
        [
            ["前端地址", "http://127.0.0.1:5173/"],
            ["后端地址", "http://127.0.0.1:8000"],
            ["登录账号", "ops001"],
            ["功能入口", "左侧导航栏：AI偏好识别"],
        ],
        widths=[1.5, 5.0],
    )

    add_para(doc, "8.2 模型权重准备", "Heading 2")
    add_para(
        doc,
        "当前环境已通过 ModelScope 下载 BGE 模型权重。如果重新部署到新机器，需要先安装依赖并下载模型。"
        "依赖包括 torch、transformers、sentence-transformers、modelscope。"
    )
    add_code(
        doc,
        r"""
cd E:\shixun\zy
python -m venv backend\.venv
.\backend\.venv\Scripts\python.exe -m pip install -e backend
.\backend\.venv\Scripts\python.exe -m pip install modelscope

# 下载模型权重
.\backend\.venv\Scripts\python.exe backend\scripts\download_ai_model.py
        """,
    )

    add_para(doc, "9. 测试与验证", "Heading 1")
    add_para(doc, "9.1 已完成验证", "Heading 2")
    add_bullet(doc, "前端构建验证：npm run build 通过。")
    add_bullet(doc, "AI 服务单元测试：test_ai_preference_service.py 通过。")
    add_bullet(doc, "接口验证：/api/ai/preference-categories 返回 model.loaded = true。")
    add_bullet(doc, "页面验证：AI偏好识别页面显示“已加载”“语义向量”“BAAI/bge-small-zh-v1.5”。")
    add_bullet(doc, "演示结果验证：user001 识别为鞋服，user002 识别为食品，user003 识别为数码。")

    add_para(doc, "9.2 验证命令", "Heading 2")
    add_code(
        doc,
        r"""
cd E:\shixun\zy
pytest backend\tests\unit\test_ai_preference_service.py -q

cd E:\shixun\zy\frontend
npm run build
        """,
    )

    add_para(doc, "10. 边界与后续优化", "Heading 1")
    add_para(doc, "10.1 当前边界", "Heading 2")
    add_bullet(doc, "当前模块使用预训练模型做向量表示，没有针对本项目数据重新训练模型。")
    add_bullet(doc, "偏好分类结果依赖商品 metadata 的完整性，商品名称、分类、品牌、标签越完整，结果越稳定。")
    add_bullet(doc, "当前以用户行为事件为主要输入，暂未引入用户人口属性、历史订单长期画像等更复杂特征。")
    add_bullet(doc, "当前演示环境建议使用内存仓库，生产环境可接 MySQL 持久化行为事件。")

    add_para(doc, "10.2 后续优化方向", "Heading 2")
    add_bullet(doc, "引入更多商品文本字段，如详情页卖点、用户评价摘要、搜索词上下文。")
    add_bullet(doc, "将用户兴趣向量持久化，减少重复计算，提高查询性能。")
    add_bullet(doc, "增加时间窗口筛选，例如近 7 天偏好、近 30 天偏好、长期偏好。")
    add_bullet(doc, "增加推荐解释，例如“因为用户近期支付了跑鞋并加购了板鞋，所以判定鞋服偏好较高”。")
    add_bullet(doc, "后续可用真实成交数据微调品类分类器，使模型更贴近业务场景。")

    add_para(doc, "11. 答辩说明建议", "Heading 1")
    add_callout(
        doc,
        "推荐表述",
        "本系统使用 BGE 中文语义向量模型对商品文本进行向量化表示。模型本身不直接输出用户偏好类别，"
        "而是提供商品语义理解能力；系统在此基础上结合用户浏览、点击、收藏、加购、支付等行为权重，"
        "聚合形成用户兴趣向量，并通过相似度计算识别用户偏好的商品品类。",
        fill="E8F4FF",
    )
    add_para(
        doc,
        "如果被问“这个 AI 到底做了什么”，可以回答：AI 模型负责理解商品文本含义，把文字转成语义向量；"
        "业务算法负责根据用户行为把这些向量聚合起来，并判断用户更偏好哪类商品。"
    )

    add_para(doc, "12. 参考资料", "Heading 1")
    add_bullet(doc, "C-Pack: Packed Resources For General Chinese Embeddings, arXiv:2309.07597, 2023.")
    add_bullet(doc, "BAAI / FlagEmbedding 项目：BGE 中文语义向量模型说明与使用文档。")
    add_bullet(doc, "Hugging Face / ModelScope 模型页：BAAI/bge-small-zh-v1.5。")
    add_bullet(doc, "本项目源码：backend/src/services/ai_preference_service.py 与 frontend/src/App.tsx。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
