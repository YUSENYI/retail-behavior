#!/usr/bin/env python3
"""将两份差异化实训报告 Markdown 转为格式化 Word 文档"""

import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def create_docx(md_path, output_path, name, student_id):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 全局样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    def add_para(text='', font_cn='宋体', font_en='Times New Roman', size=12,
                 bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, indent=True,
                 sp_before=0, sp_after=0, line_sp=1.5):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.line_spacing = line_sp
        pf.space_before = Pt(sp_before)
        pf.space_after = Pt(sp_after)
        if indent and text.strip():
            pf.first_line_indent = Cm(0.75)
        if text:
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.font.name = font_en
            r._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
            r.bold = bold
        return p

    def empty_line(n=1):
        for _ in range(n):
            add_para(sp_before=0, sp_after=0, indent=False)

    # ===== 封面 =====
    empty_line(2)
    add_para('大连民族大学', '宋体', 'Times New Roman', 22, True,
             WD_ALIGN_PARAGRAPH.CENTER, False)
    empty_line(1)
    add_para('AI驱动的企业级项目实战课程', '宋体', 'Times New Roman', 20, True,
             WD_ALIGN_PARAGRAPH.CENTER, False)
    add_para('实训报告', '宋体', 'Times New Roman', 36, True,
             WD_ALIGN_PARAGRAPH.CENTER, False)
    empty_line(2)
    add_para('设计报告题目：智能零售用户行为分析系统设计与开发',
             '宋体', 'Times New Roman', 14, True,
             WD_ALIGN_PARAGRAPH.CENTER, False, sp_before=24, sp_after=24)
    empty_line(1)
    add_para(f'学生姓名：{name}', '宋体', 'Times New Roman', 12, False,
             WD_ALIGN_PARAGRAPH.CENTER, False, sp_before=8, sp_after=8)
    add_para(f'学生学号：{student_id}', '宋体', 'Times New Roman', 12, False,
             WD_ALIGN_PARAGRAPH.CENTER, False, sp_before=8, sp_after=8)
    add_para('组队名称：', '宋体', 'Times New Roman', 12, False,
             WD_ALIGN_PARAGRAPH.CENTER, False, sp_before=8, sp_after=8)
    add_para('指导教师：', '宋体', 'Times New Roman', 12, False,
             WD_ALIGN_PARAGRAPH.CENTER, False, sp_before=8, sp_after=8)
    empty_line(2)
    add_para('日  期：2026年 06月', '宋体', 'Times New Roman', 12, False,
             WD_ALIGN_PARAGRAPH.CENTER, False, sp_before=12, sp_after=0)

    # ===== 正文 =====
    doc.add_page_break()

    # 解析 markdown 内容
    lines = content.split('\n')

    # 跳过封面部分（# 大连民族大学 到 --- 之间的内容）、目录部分
    in_body = False
    after_toc = False

    for line in lines:
        stripped = line.strip()

        # 跳过markdown封面
        if stripped == '---' and not in_body:
            continue
        if not in_body and '# 大连民族大学' in stripped:
            continue
        if not in_body and 'AI驱动的企业级项目实战课程' in stripped:
            continue
        if not in_body and '实训报告' in stripped and 'AI编程' not in stripped and len(stripped) < 10:
            continue
        if not in_body and '<br>' in stripped:
            continue
        if not in_body and '设计报告题目' in stripped:
            continue
        if not in_body and '学生姓名' in stripped:
            continue
        if not in_body and '学生学号' in stripped:
            continue
        if not in_body and '组队名称' in stripped:
            continue
        if not in_body and '指导教师' in stripped:
            continue
        if not in_body and '日' in stripped and '期' in stripped:
            continue
        if not in_body and stripped == '---':
            in_body = True
            continue

        if not in_body:
            continue

        # 报告题目区域
        if '报告题目：' in stripped and '##' in stripped:
            add_para('报告题目：', '宋体', 'Times New Roman', 12, False,
                     WD_ALIGN_PARAGRAPH.LEFT, False)
            continue
        if '智能零售用户行为分析系统设计与开发' in stripped and '##' not in stripped and '——' not in stripped:
            add_para('智能零售用户行为分析系统设计与开发', '宋体', 'Times New Roman', 14, True,
                     WD_ALIGN_PARAGRAPH.LEFT, False, sp_before=2)
            continue
        if '基于AI编程的企业级项目实战' in stripped:
            add_para('——基于AI编程的企业级项目实战', '宋体', 'Times New Roman', 12, False,
                     WD_ALIGN_PARAGRAPH.LEFT, False, sp_before=2, sp_after=16)
            continue

        # 摘要
        if stripped.startswith('**摘要：**'):
            text = stripped.replace('**摘要：**', '')
            add_para('摘要：' + text, '宋体', 'Times New Roman', 12, False,
                     WD_ALIGN_PARAGRAPH.LEFT, True)
            continue
        if stripped.startswith('本文') and ('介绍' in stripped or '详细阐述' in stripped):
            add_para(stripped, '宋体', 'Times New Roman', 12, False,
                     WD_ALIGN_PARAGRAPH.LEFT, True)
            continue
        if stripped.startswith('**关键词：**'):
            text = stripped.replace('**关键词：**', '')
            add_para('关键词：' + text, '宋体', 'Times New Roman', 12, False,
                     WD_ALIGN_PARAGRAPH.LEFT, True, sp_before=4)
            continue

        # 目录
        if '目' in stripped and '录' in stripped and stripped.replace(' ', '').replace('&emsp;', '') == '目录':
            empty_line(2)
            add_para('目  录', '黑体', 'Times New Roman', 18, True,
                     WD_ALIGN_PARAGRAPH.CENTER, False, sp_before=12, sp_after=16)
            after_toc = True
            continue

        if after_toc and stripped.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.')):
            toc_items = ['1 引言', '2 系统整体设计/需求分析', '3 功能模块设计',
                         '4 主要流程设计', '5 界面设计图', '6 课程总结', '7 参考文献']
            for item in toc_items:
                add_para(item, '宋体', 'Times New Roman', 12, False,
                         WD_ALIGN_PARAGRAPH.LEFT, False, sp_before=4, sp_after=4)
            empty_line(4)
            doc.add_page_break()
            after_toc = False
            continue

        # 章节标题
        section_match = re.match(r'^## (\d+ .+)', stripped)
        if section_match:
            if after_toc:
                continue
            title = section_match.group(1)
            add_para(title, '黑体', 'Times New Roman', 16, True,
                     WD_ALIGN_PARAGRAPH.LEFT, False,
                     sp_before=18 if title[0] != '1' else 12, sp_after=6)
            continue

        # 子章节
        sub_match = re.match(r'^### (\d+\.\d+ .+)', stripped)
        if sub_match:
            add_para(sub_match.group(1), '黑体', 'Times New Roman', 14, True,
                     WD_ALIGN_PARAGRAPH.LEFT, False, sp_before=12, sp_after=4)
            continue

        # 空行
        if stripped == '<br>' or stripped == '':
            continue
        if stripped == '---':
            continue

        # 表格行
        if '|' in stripped and not stripped.startswith('```'):
            continue

        # 代码块
        if stripped == '```':
            continue

        # 架构图 ascii
        if stripped.startswith('┌') or stripped.startswith('│') or stripped.startswith('└'):
            continue

        # 图片标记位
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            if text:
                add_para('图片标记位：' + text, '宋体', 'Times New Roman', 12, True,
                         WD_ALIGN_PARAGRAPH.CENTER, False, sp_before=6, sp_after=6)
            continue

        # 参考文献条目
        if re.match(r'^\[\d+\]', stripped):
            add_para(stripped, '宋体', 'Times New Roman', 12, False,
                     WD_ALIGN_PARAGRAPH.LEFT, True, sp_before=2, sp_after=2)
            continue

        # 带首行缩进的编号段落（列表项）
        if re.match(r'^[（(]\d+[）)]', stripped) or re.match(r'^\d+[\.\、]', stripped):
            add_para(stripped, '宋体', 'Times New Roman', 12, False,
                     WD_ALIGN_PARAGRAPH.LEFT, True, sp_before=2, sp_after=2)
            continue

        # 普通段落
        if stripped:
            add_para(stripped, '宋体', 'Times New Roman', 12, False,
                     WD_ALIGN_PARAGRAPH.LEFT, True, sp_before=2, sp_after=2)

    # 封底
    empty_line(4)
    add_para('AI编程课程—实训报告', '宋体', 'Times New Roman', 12, False,
             WD_ALIGN_PARAGRAPH.CENTER, False, sp_before=12)

    doc.save(output_path)
    print(f'OK: {output_path}')


if __name__ == '__main__':
    base = 'E:/shixun/zy/文档'
    create_docx(f'{base}/王正滔_智能零售用户行为分析系统_实训报告.md',
                f'{base}/王正滔_智能零售用户行为分析系统_实训报告.docx',
                '王正滔', '11')
    create_docx(f'{base}/戴天笑_智能零售用户行为分析系统_实训报告.md',
                f'{base}/戴天笑_智能零售用户行为分析系统_实训报告.docx',
                '戴天笑', '6')
