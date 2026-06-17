"""分析原版 docx 结构"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

doc = Document(r"E:\shixun\zy\文档\尤照程_智能零售用户行为分析系统_实训报告.docx")

print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")
print(f"节数: {len(doc.sections)}")

# 检查 sections
for i, section in enumerate(doc.sections):
    print(f"\nSection {i}:")
    print(f"  page_width: {section.page_width}")
    print(f"  page_height: {section.page_height}")
    print(f"  left_margin: {section.left_margin}")
    print(f"  right_margin: {section.right_margin}")
    print(f"  top_margin: {section.top_margin}")
    print(f"  bottom_margin: {section.bottom_margin}")

# 检查图片
from docx.opc.constants import RELATIONSHIP_TYPE as RT
image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image_count += 1
        print(f"\nImage {image_count}: reltype={rel.reltype}, target_ref={rel.target_ref}")

print(f"\n总图片数: {image_count}")

# 列出每段的关键信息
print("\n========== 段落详情 ==========")
for i, para in enumerate(doc.paragraphs):
    # 检查是否有图片
    has_image = False
    for run in para.runs:
        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            has_image = True
        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
            has_image = True
        # Check for inline shapes
        drawings = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
        if drawings:
            has_image = True
    
    text_preview = para.text[:80] if para.text else "(空)"
    style = para.style.name if para.style else "None"
    alignment = para.alignment
    fmt = para.paragraph_format
    print(f"\n--- Para {i} (style={style}, align={alignment}) ---")
    print(f"  text: {text_preview}")
    if has_image:
        print(f"  [CONTAINS IMAGE]")
    
    # 每个 run 的格式
    for j, run in enumerate(para.runs):
        font = run.font
        r_text = run.text[:50] if run.text else "(无文字)"
        print(f"  Run {j}: font={font.name}, size={font.size}, bold={font.bold}, color={font.color.rgb if font.color and font.color.rgb else 'N/A'}, text='{r_text}'")

print("\n========== 表格详情 ==========")
for t_idx, table in enumerate(doc.tables):
    print(f"\n表格 {t_idx}: {len(table.rows)}行 x {len(table.columns)}列")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text[:60]
            if text.strip():
                print(f"  [{r_idx},{c_idx}]: {text}")
